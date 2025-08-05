import os
import zipfile
import pandas as pd
import glob
import json
from openai import OpenAI
import subprocess
import docx  # 用于处理 .docx 文件
import PyPDF2 # 用于处理 .pdf 文件
from PIL import Image # 用于处理图片
import pytesseract     # 用于OCR
import io              # 用于处理内存中的字节流
# --- 配置 ---
# 脚本将自动从环境变量中读取 'OPENAI_API_KEY' 和 'OPENAI_BASE_URL'。
if "OPENAI_API_KEY" not in os.environ:
    print("错误：未找到 OpenAI API 密钥。")
    print("请设置环境变量 'OPENAI_API_KEY'。")
    exit()

client = OpenAI()
LLM_MODEL = "deepseek-chat"

def extract_text_from_folder(folder_path):
    """
    (已修改) 递归地从多种文件类型中提取文本。
    """
    print(f"  - 开始从 {os.path.basename(folder_path)} 提取文本...")
    full_text = ""
    # 支持的扩展名列表
    supported_extensions = ['*.txt', '*.md', '*.py', '*.java', '*.pdf', '*.docx', '*.doc']
    
    files_to_process = []
    for ext in supported_extensions:
        # 递归查找所有匹配的文件
        files_to_process.extend(glob.glob(os.path.join(folder_path, '**', ext), recursive=True))

    if not files_to_process:
        print("    - 未找到支持的文本文件。")
        return "[内容为空或文件格式不支持]"

    for file_path in files_to_process:
        content = ""
        file_name = os.path.basename(file_path)
        file_ext = os.path.splitext(file_name)[1].lower()
        
        try:
            if file_ext == '.docx':
                print(f"    - 正在处理 DOCX 文件: {file_name} (包含OCR)")
                doc = docx.Document(file_path)
                content_parts = []
                # 遍历文档中的所有段落
                for para in doc.paragraphs:
                    # 1. 添加段落本身的文本
                    if para.text.strip():
                        content_parts.append(para.text)
                    
                    # 2. 检查段落内是否包含图片并进行OCR
                    # 图片存在于段落的 "run" 元素中
                    for run in para.runs:
                        # run.element.xpath('.//pic:pic') 查找内联图片
                        if run.element.xpath('.//pic:pic'): 
                            # 获取图片的 rId
                            for r in run.element.xpath(".//a:blip"):
                                rId = r.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                                if rId:
                                    try:
                                        image_part = doc.part.related_parts[rId]
                                        image_stream = io.BytesIO(image_part.blob)
                                        image = Image.open(image_stream)
                                        
                                        # 使用Tesseract进行OCR，同时支持英文和简体中文
                                        ocr_text = pytesseract.image_to_string(image, lang='eng+chi_sim')
                                        
                                        if ocr_text.strip():
                                            # 将OCR结果作为一个独立的部分添加
                                            ocr_block = f"\n--- [图片OCR内容开始] ---\n{ocr_text.strip()}\n--- [图片OCR内容结束] ---\n"
                                            content_parts.append(ocr_block)
                                            print(f"      - 成功对 {file_name} 中的一张图片进行OCR。")
                                    except Exception as ocr_error:
                                        print(f"      - 警告: 对 {file_name} 中的图片进行OCR时失败: {ocr_error}")
                
                content = "\n".join(content_parts)
            elif file_ext == '.pdf':
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    # 检查PDF是否加密
                    if reader.is_encrypted:
                        print(f"    - 警告: 跳过加密的PDF文件: {file_name}")
                        content = "[文件已加密，无法提取内容]"
                    else:
                        for page in reader.pages:
                            extracted = page.extract_text()
                            if extracted:
                                content += extracted + "\n"
            # elif file_ext == '.doc':
            #     # 尝试使用 libreoffice 将 .doc 转换为 .docx 进行处理
            #     # 这需要系统上安装了 LibreOffice
            #     docx_path = os.path.splitext(file_path)[0] + '.docx'
            #     try:
            #         # 检查转换后的文件是否已存在，如果存在则跳过转换
            #         if not os.path.exists(docx_path):
            #             print(f"    - 正在尝试将 {file_name} 转换为 .docx...")
            #             # 使用 subprocess 调用 libreoffice 进行转换
            #             subprocess.run(
            #                 ['libreoffice', '--headless', '--convert-to', 'docx', '--outdir', os.path.dirname(file_path), file_path],
            #                 check=True,
            #                 stdout=subprocess.DEVNULL,
            #                 stderr=subprocess.DEVNULL
            #             )
                    
            #         # 读取转换后的 .docx 文件
            #         if os.path.exists(docx_path):
            #             print(f"    - 已从转换后的 {os.path.basename(docx_path)} 中读取内容。")
            #             doc = docx.Document(docx_path)
            #             content = "\n".join([para.text for para in doc.paragraphs])
            #             # (可选) 清理转换生成的文件
            #             # os.remove(docx_path)
            #         else:
            #             raise FileNotFoundError("转换后的 .docx 文件未找到。")

            #     except (FileNotFoundError, subprocess.CalledProcessError) as e:
            #         # FileNotFoundError: libreoffice 未安装或不在 PATH 中
            #         # CalledProcessError: libreoffice 转换失败
            #         print(f"    - 警告: 无法处理 .doc 文件: {file_name}。请确保已安装 LibreOffice。错误: {e}")
            #         content = "[.doc 文件处理失败，请确保已安装 LibreOffice 或手动转换为 .docx]"
            elif file_ext in ['.txt', '.md', '.py', '.java']:
                # 处理常见的纯文本文件
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
            
            if content:
                full_text += f"\n\n--- 文件: {file_name} ---\n\n{content}"

        except Exception as e:
            error_message = f"提取失败: {e}"
            full_text += f"\n\n--- 文件: {file_name} ({error_message}) ---\n\n"
            print(f"    - 警告: 无法从 {file_path} 提取文本: {e}")
            
    if not full_text:
        return "[内容为空或所有文件均无法提取]"
        
    return full_text

def analyze_with_llm(student_content, rubric):
    """
    将学生内容和评分标准发送给 LLM 进行评分，并返回分数和评语。
    """
    system_prompt = """
    你是一名经验丰富的大学计算机课程助教，你的任务是根据提供的评分标准，对学生的实验报告和代码进行细致、公正的评分。
    请严格按照评分标准中的每一个要点进行分析。
    你的输出必须是一个JSON对象，包含两个键：
    1. 'score' (一个浮点数): 最终得分。
    2. 'comment' (一个字符串): 评分评语。一般情况是4分，如果分数低于4分，评语必须清晰地指出显著缺陷和扣分原因。如果得分是4分或更高，评语可以是“表现良好”或类似的鼓励性话语。
    """

    user_prompt = f"""
    请根据以下【评分标准】对这位学生的【提交内容】进行评分。

    【评分标准】
    {rubric}

    【学生提交内容】
    {student_content}

    请严格按照评分标准进行分析，并以JSON格式返回评分结果，包含 'score' 和 'comment' 两个键。
    """

    try:
        print("  - 正在调用 LLM 进行分析...")
        response = client.chat.completions.create(
            model=LLM_MODEL,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            response_format={"type": "json_object"},
            temperature=0.2,
        )
        
        result_json = json.loads(response.choices[0].message.content)
        score = float(result_json.get('score', 2.0))
        comment = result_json.get('comment', "LLM未提供评语，请手动检查。")
        print(f"  - LLM分析完成。")
        return score, comment

    except Exception as e:
        print(f"  - LLM API 调用失败: {e}")
        return 2.0, f"LLM分析失败，请手动评分。错误信息: {e}"


def main(current_dir):
    """
    主函数，遍历学生文件夹，处理内部的zip文件，使用LLM分析，并创建Excel报告。
    """
    try:
        # 假设评分标准在脚本所在目录或工作目录
        rubric_path = os.path.join(os.path.dirname(__file__), '测试管理评分标准.txt')
        if not os.path.exists(rubric_path):
            rubric_path = '测试管理评分标准.txt' # Fallback to current working directory
        with open(rubric_path, 'r', encoding='utf-8') as f:
            rubric = f.read()
    except FileNotFoundError:
        print("错误: '评分标准.txt' 文件未找到。请确保它与脚本在同一目录或在您运行脚本的目录中。")
        return

    # 获取所有项目，并筛选出非'scripts'的文件夹
    all_items = os.listdir(current_dir)
    student_folders = [
        item for item in all_items
        if os.path.isdir(os.path.join(current_dir, item)) and item != 'scripts'
    ]

    if not student_folders:
        print(f"在 '{current_dir}' 目录下未找到任何学生作业文件夹。")
        return

    results = []
    print(f"找到 {len(student_folders)} 个学生文件夹，开始处理...")

    for student_name in student_folders:
        extract_path = os.path.join(current_dir, student_name)
        print(f"\n正在处理: {student_name}")

        # 递归查找并解压文件夹内的所有 .zip 文件
        try:
            nested_zips = glob.glob(os.path.join(extract_path, '**', '*.zip'), recursive=True)
            if nested_zips:
                print(f"  - 在 {student_name} 中找到 {len(nested_zips)} 个zip文件，正在解压...")
                for zip_file_path in nested_zips:
                    zip_dir = os.path.dirname(zip_file_path)
                    with zipfile.ZipFile(zip_file_path, 'r') as zf:
                        zf.extractall(zip_dir)
                    print(f"    - 已解压: {os.path.relpath(zip_file_path, current_dir)}")
                    # 安全起见，可以选择删除已解压的zip文件
                    # os.remove(zip_file_path)
            else:
                print("  - 未在文件夹内找到需要解压的zip文件。")
        except zipfile.BadZipFile as e:
            print(f"  - 错误: {e}。跳过此学生。")
            score, comment = 2.0, f"文件夹中包含损坏的ZIP文件: {e}"
            results.append({'学生/文件名': student_name, '分数': score, '评语': comment})
            continue
        except Exception as e:
            print(f"  - 解压过程中发生未知错误: {e}。跳过此学生。")
            score, comment = 2.0, f"解压时发生错误: {e}"
            results.append({'学生/文件名': student_name, '分数': score, '评语': comment})
            continue

        student_content = extract_text_from_folder(extract_path)
        
        score, comment = analyze_with_llm(student_content, rubric)
        results.append({'学生/文件名': student_name, '分数': score, '评语': comment})
        print(f"  - 评分完成: {student_name}, 得分: {score}")

    df = pd.DataFrame(results)
    output_filename = os.path.join(current_dir, 'LLM_评分结果.xlsx')
    df.to_excel(output_filename, index=False, engine='openpyxl')
    print(f"\n所有评分完成！结果已保存到 {output_filename}")

if __name__ == '__main__':
    main('/home/tyrfly/labtask')