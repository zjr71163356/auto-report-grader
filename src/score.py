import PyPDF2
import io              # 用于处理内存中的字节流
import pytesseract     # 用于OCR
from PIL import Image  # 用于处理图片
import os
import zipfile
import pandas as pd
import glob
import json
from openai import OpenAI
import subprocess
import docx  # 用于处理 .docx 文件

# 导入配置
from config import (
    OPENAI_API_KEY, OPENAI_BASE_URL, LLM_MODEL,
    COLLECTED_DIR, RUBRIC_FILE, OUTPUT_FILENAME,
    SUPPORTED_EXTENSIONS, ZIP_EXTENSIONS, RAR_EXTENSIONS,
    OCR_LANGUAGES, SCORING_TEMPERATURE,
    MIN_SCORE, MAX_SCORE, DEFAULT_SCORE,
    VERBOSE_LOGGING
)

try:
    import rarfile
    RARFILE_AVAILABLE = True
except ImportError:
    RARFILE_AVAILABLE = False
    print("警告: rarfile 包未安装，无法解压 RAR 文件。请运行 'pip install rarfile' 安装。")
# --- 配置 ---
# API 配置检查
if not OPENAI_API_KEY:
    print("错误：未找到 OpenAI API 密钥。")
    print("请设置环境变量 'OPENAI_API_KEY'。")
    exit()

# 初始化 OpenAI 客户端
client = OpenAI(
    api_key=OPENAI_API_KEY,
    base_url=OPENAI_BASE_URL
)


def extract_text_from_folder(folder_path):
    """
    (已修改) 递归地从多种文件类型中提取文本。
    """
    if VERBOSE_LOGGING:
        print(f"  - 开始从 {os.path.basename(folder_path)} 提取文本...")
    full_text = ""
    # 使用配置中的支持扩展名列表
    supported_extensions = SUPPORTED_EXTENSIONS

    files_to_process = []
    for ext in supported_extensions:
        # 递归查找所有匹配的文件
        files_to_process.extend(glob.glob(os.path.join(
            folder_path, '**', ext), recursive=True))

    if not files_to_process:
        if VERBOSE_LOGGING:
            print("    - 未找到支持的文本文件。")
        return "[内容为空或文件格式不支持]"

    for file_path in files_to_process:
        content = ""
        file_name = os.path.basename(file_path)
        file_ext = os.path.splitext(file_name)[1].lower()

        try:
            if file_ext == '.docx':
                print(f"    - 正在处理 DOCX 文件: {file_name} (包含OCR)")
                doc = docx.Document(file_path)
                content_parts = []
                # 遍历文档中的所有段落
                for para in doc.paragraphs:
                    # 1. 添加段落本身的文本
                    if para.text.strip():
                        content_parts.append(para.text)

                    # 2. 检查段落内是否包含图片并进行OCR
                    # 图片存在于段落的 "run" 元素中
                    for run in para.runs:
                        # run.element.xpath('.//pic:pic') 查找内联图片
                        if run.element.xpath('.//pic:pic'):
                            # 获取图片的 rId
                            for r in run.element.xpath(".//a:blip"):
                                rId = r.get(
                                    '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                                if rId:
                                    try:
                                        image_part = doc.part.related_parts[rId]
                                        image_stream = io.BytesIO(
                                            image_part.blob)
                                        image = Image.open(image_stream)

                                        # 使用配置中的OCR语言设置
                                        ocr_text = pytesseract.image_to_string(
                                            image, lang=OCR_LANGUAGES)

                                        if ocr_text.strip():
                                            # 将OCR结果作为一个独立的部分添加
                                            ocr_block = f"\n--- [图片OCR内容开始] ---\n{ocr_text.strip()}\n--- [图片OCR内容结束] ---\n"
                                            content_parts.append(ocr_block)
                                            print(
                                                f"      - 成功对 {file_name} 中的一张图片进行OCR。")
                                    except Exception as ocr_error:
                                        print(
                                            f"      - 警告: 对 {file_name} 中的图片进行OCR时失败: {ocr_error}")

                content = "\n".join(content_parts)
            elif file_ext == '.pdf':
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    # 检查PDF是否加密
                    if reader.is_encrypted:
                        print(f"    - 警告: 跳过加密的PDF文件: {file_name}")
                        content = "[文件已加密，无法提取内容]"
                    else:
                        for page in reader.pages:
                            extracted = page.extract_text()
                            if extracted:
                                content += extracted + "\n"
            elif file_ext == '.doc':
                # 尝试使用 libreoffice 将 .doc 转换为 .docx 进行处理
                # 这需要系统上安装了 LibreOffice
                docx_path = os.path.splitext(file_path)[0] + '.docx'
                try:
                    # 检查转换后的文件是否已存在，如果存在则跳过转换
                    if not os.path.exists(docx_path):
                        print(f"    - 正在尝试将 {file_name} 转换为 .docx...")
                        # 使用 subprocess 调用 libreoffice 进行转换
                        subprocess.run(
                            ['libreoffice', '--headless', '--convert-to', 'docx',
                                '--outdir', os.path.dirname(file_path), file_path],
                            check=True,
                            stdout=subprocess.DEVNULL,
                            stderr=subprocess.DEVNULL
                        )

                    # 读取转换后的 .docx 文件
                    if os.path.exists(docx_path):
                        print(
                            f"    - 已从转换后的 {os.path.basename(docx_path)} 中读取内容。")
                        doc = docx.Document(docx_path)
                        content = "\n".join(
                            [para.text for para in doc.paragraphs])
                        # (可选) 清理转换生成的文件
                        # os.remove(docx_path)
                    else:
                        raise FileNotFoundError("转换后的 .docx 文件未找到。")

                except (FileNotFoundError, subprocess.CalledProcessError) as e:
                    # FileNotFoundError: libreoffice 未安装或不在 PATH 中
                    # CalledProcessError: libreoffice 转换失败
                    print(
                        f"    - 警告: 无法处理 .doc 文件: {file_name}。请确保已安装 LibreOffice。错误: {e}")
                    content = "[.doc 文件处理失败，请确保已安装 LibreOffice 或手动转换为 .docx]"
            elif file_ext in ['.txt', '.md', '.py', '.java']:
                # 处理常见的纯文本文件
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()

            if content:
                full_text += f"\n\n--- 文件: {file_name} ---\n\n{content}"

        except Exception as e:
            error_message = f"提取失败: {e}"
            full_text += f"\n\n--- 文件: {file_name} ({error_message}) ---\n\n"
            print(f"    - 警告: 无法从 {file_path} 提取文本: {e}")

    if not full_text:
        return "[内容为空或所有文件均无法提取]"

    return full_text


def analyze_with_llm(student_content, rubric):
    """
    将学生内容和评分标准发送给 LLM 进行评分，并返回分数和评语。
    """
    system_prompt = """
    你是一名经验丰富的大学计算机课程助教，你的任务是根据提供的评分标准，对学生的软件测试综合实验报告进行细致、公正的评分。

    重要评分原则：
    1. 严格按照评分标准的10分制进行评分
    2. 特别注意识别疑似大模型直接生成的报告（markdown风格明显、图表过于规整等），如未提供prompt过程则不能评为9-10分
    3. 重视报告的个性化、分析思考深度和实验覆盖度
    4. 考虑比例控制：10分(10-15%)、9分(15-25%)、6分及以下(10-20%)，平均分应在8分左右

    你的输出必须是一个JSON对象，包含两个键：
    1. 'score' (一个浮点数): 最终得分(0-10分)
    2. 'comment' (一个字符串): 详细的评分评语，说明得分理由和改进建议
    """

    user_prompt = f"""
    请根据以下【评分标准】对这位学生的【软件测试综合实验报告】进行评分。

    【评分标准】
    {rubric}

    【学生提交内容】
    {student_content}

    请仔细分析报告的以下方面：
    1. 报告格式的清晰性和规范性
    2. 实验要求的覆盖程度
    3. 个性化内容的体现
    4. 分析与思考的深度
    5. 是否疑似大模型直接生成（注意格式风格和是否提供prompt过程）

    根据10分制评分标准，以JSON格式返回评分结果，包含 'score' 和 'comment' 两个键。
    """

    try:
        if VERBOSE_LOGGING:
            print("  - 正在调用 LLM 进行分析...")
        response = client.chat.completions.create(
            model=LLM_MODEL,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            response_format={"type": "json_object"},
            temperature=SCORING_TEMPERATURE,
        )

        result_json = json.loads(response.choices[0].message.content)
        score = float(result_json.get('score', DEFAULT_SCORE))
        comment = result_json.get('comment', "LLM未提供评语，请手动检查。")

        # 确保分数在配置的合理范围内
        score = max(MIN_SCORE, min(MAX_SCORE, score))

        if VERBOSE_LOGGING:
            print(f"  - LLM分析完成。")
        return score, comment

    except Exception as e:
        print(f"  - LLM API 调用失败: {e}")
        return DEFAULT_SCORE, f"LLM分析失败，请手动评分。错误信息: {e}"


def main(current_dir=None, rubric_path=None):
    """
    主函数，遍历学生文件夹，处理内部的zip文件，使用LLM分析，并创建Excel报告。
    """
    # 使用配置中的默认路径，如果没有提供参数
    current_dir = current_dir or str(COLLECTED_DIR)
    rubric_path = rubric_path or str(RUBRIC_FILE)

    try:
        # 读取评分标准文件
        with open(rubric_path, 'r', encoding='utf-8') as f:
            rubric = f.read()
    except FileNotFoundError:
        print(f"错误: 评分标准文件 '{rubric_path}' 未找到。")
        return

    # 获取所有学生文件夹
    all_items = os.listdir(current_dir)
    student_folders = [
        item for item in all_items
        if os.path.isdir(os.path.join(current_dir, item))
    ]

    if not student_folders:
        print(f"在 '{current_dir}' 目录下未找到任何学生作业文件夹。")
        return

    results = []
    print(f"找到 {len(student_folders)} 个学生文件夹，开始处理...")

    for student_name in student_folders:
        extract_path = os.path.join(current_dir, student_name)
        print(f"\n正在处理: {student_name}")

        # 递归查找并解压文件夹内的所有压缩文件
        try:
            # 查找压缩文件
            zip_files = []
            for ext in ZIP_EXTENSIONS:
                zip_files.extend(glob.glob(os.path.join(
                    extract_path, '**', ext), recursive=True))

            rar_files = []
            for ext in RAR_EXTENSIONS:
                rar_files.extend(glob.glob(os.path.join(
                    extract_path, '**', ext), recursive=True))

            # 解压 ZIP 文件
            if zip_files:
                print(
                    f"  - 在 {student_name} 中找到 {len(zip_files)} 个zip文件，正在解压...")
                for zip_file_path in zip_files:
                    try:
                        zip_dir = os.path.dirname(zip_file_path)
                        with zipfile.ZipFile(zip_file_path, 'r') as zf:
                            zf.extractall(zip_dir)
                        print(
                            f"    - 已解压: {os.path.relpath(zip_file_path, current_dir)}")
                        # 解压成功后删除原压缩文件
                        os.remove(zip_file_path)
                        print(
                            f"    - 已删除原文件: {os.path.basename(zip_file_path)}")
                    except zipfile.BadZipFile:
                        print(
                            f"    - 警告: {os.path.basename(zip_file_path)} 不是有效的ZIP文件")
                    except Exception as e:
                        print(
                            f"    - 警告: 解压 {os.path.basename(zip_file_path)} 失败: {e}")

            # 解压 RAR 文件
            if rar_files:
                if RARFILE_AVAILABLE:
                    print(
                        f"  - 在 {student_name} 中找到 {len(rar_files)} 个rar文件，正在解压...")
                    for rar_file_path in rar_files:
                        try:
                            rar_dir = os.path.dirname(rar_file_path)
                            with rarfile.RarFile(rar_file_path) as rf:
                                rf.extractall(rar_dir)
                            print(
                                f"    - 已解压: {os.path.relpath(rar_file_path, current_dir)}")
                            # 解压成功后删除原压缩文件
                            os.remove(rar_file_path)
                            print(
                                f"    - 已删除原文件: {os.path.basename(rar_file_path)}")
                        except rarfile.RarCannotExec:
                            print(
                                f"    - 警告: 无法解压 {os.path.basename(rar_file_path)}，请确保系统已安装 unrar 工具")
                        except Exception as e:
                            print(
                                f"    - 警告: 解压 {os.path.basename(rar_file_path)} 失败: {e}")
                else:
                    print(
                        f"  - 发现 {len(rar_files)} 个RAR文件，但 rarfile 包未安装，请运行 'pip install rarfile' 并确保系统已安装 unrar 工具")

        except Exception as e:
            print(f"  - 查找压缩文件时发生错误: {e}")

        # 提取文本内容
        student_content = extract_text_from_folder(extract_path)

        # 解析学号和姓名
        if '_' in student_name:
            parts = student_name.split('_', 1)
            student_id = parts[0]
            student_real_name = parts[1]
        else:
            student_id = student_name
            student_real_name = student_name

        # 使用 LLM 进行评分
        score, comment = analyze_with_llm(student_content, rubric)

        results.append({
            '学号': student_id,
            '姓名': student_real_name,
            '文件夹名': student_name,
            '分数': score,
            '评语': comment
        })
        print(f"  - 评分完成: {student_name}, 得分: {score}")

    # 生成 Excel 报告
    df = pd.DataFrame(results)
    output_filename = os.path.join(current_dir, OUTPUT_FILENAME)
    df.to_excel(output_filename, index=False, engine='openpyxl')
    print(f"\n所有评分完成！结果已保存到 {output_filename}")

    # 显示评分统计
    print("\n评分统计:")
    print(f"平均分: {df['分数'].mean():.2f}")
    print(f"最高分: {df['分数'].max():.2f}")
    print(f"最低分: {df['分数'].min():.2f}")

    # 分数分布
    score_distribution = df['分数'].value_counts().sort_index()
    print("\n分数分布:")
    for score, count in score_distribution.items():
        print(f"  {score}分: {count}人")


if __name__ == '__main__':
    # 使用配置文件中的路径
    main()
